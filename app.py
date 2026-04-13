import os
import tempfile
import pandas as pd
import pytesseract
from PIL import Image
import docx
import streamlit as st
from langchain.chains import RetrievalQA
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain.text_splitter import CharacterTextSplitter
from langchain_community.vectorstores import Chroma
from pypdf.errors import PdfReadError
from openai import AuthenticationError

# Adicionar a imagem no cabeçalho
image_url = "https://cienciadosdados.com/images/CINCIA_DOS_DADOS_4.png"
st.image(image_url, use_container_width=True)

# Adicionar o nome do aplicativo
st.subheader("Q&A com IA - PLN usando LangChain")

# Componentes interativos
file_input = st.file_uploader("Upload a file", type=['pdf', 'txt', 'csv', 'docx', 'jpeg', 'png'])
openaikey = st.text_input("Enter your OpenAI API Key", type='password')
prompt = st.text_area("Enter your questions", height=160)
run_button = st.button("Run!")

select_k = st.slider("Number of relevant chunks", min_value=1, max_value=5, value=2)
select_chain_type = st.radio("Chain type", ['stuff', 'map_reduce', "refine", "map_rerank"])


def load_document(file_path, file_type):
    if file_type == 'application/pdf':
        loader = PyPDFLoader(file_path)
        return loader.load()
    elif file_type == 'text/plain':
        loader = TextLoader(file_path)
        return loader.load()
    elif file_type == 'text/csv':
        df = pd.read_csv(file_path)
        from langchain.schema import Document
        return [Document(page_content=df.to_string())]
    elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        doc = docx.Document(file_path)
        full_text = [para.text for para in doc.paragraphs]
        from langchain.schema import Document
        return [Document(page_content="\n".join(full_text))]
    elif file_type in ['image/jpeg', 'image/png']:
        text = pytesseract.image_to_string(Image.open(file_path))
        from langchain.schema import Document
        return [Document(page_content=text)]
    else:
        st.error("Unsupported file type.")
        return None


def qa(file_path, file_type, query, chain_type, k):
    try:
        documents = load_document(file_path, file_type)
        if not documents:
            return None

        text_splitter = CharacterTextSplitter(chunk_size=1000, chunk_overlap=0)
        texts = text_splitter.split_documents(documents)

        embeddings = OpenAIEmbeddings()
        db = Chroma.from_documents(texts, embeddings)
        retriever = db.as_retriever(search_type="similarity", search_kwargs={"k": k})

        chain = RetrievalQA.from_chain_type(
            llm=ChatOpenAI(model="gpt-4"),
            chain_type=chain_type,
            retriever=retriever,
            return_source_documents=True
        )
        result = chain.invoke({"query": query})
        return result
    except PdfReadError as e:
        st.error(f"Error reading PDF file: {e}")
        return None
    except AuthenticationError as e:
        st.error(f"Authentication error: {e}")
        return None
    except Exception as e:
        st.error(f"Error: {e}")
        return None


def display_result(result):
    if result:
        st.markdown("### Result:")
        st.write(result["result"])
        st.markdown("### Relevant source text:")
        for doc in result["source_documents"]:
            st.markdown("---")
            st.markdown(doc.page_content)


if run_button and file_input and openaikey and prompt:
    with st.spinner("Running QA..."):
        temp_file_path = os.path.join(tempfile.gettempdir(), file_input.name)
        with open(temp_file_path, "wb") as temp_file:
            temp_file.write(file_input.read())

        os.environ["OPENAI_API_KEY"] = openaikey

        try:
            embeddings = OpenAIEmbeddings()
            embeddings.embed_documents(["test"])
        except AuthenticationError as e:
            st.error(f"Invalid OpenAI API Key: {e}")
        else:
            result = qa(temp_file_path, file_input.type, prompt, select_chain_type, select_k)
            display_result(result)